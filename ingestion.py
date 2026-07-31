"""
AutoDS-Agent :: Ingestion Engine
================================

Turns *anything* the user throws at the app into a dictionary of clean,
analysis-ready ``pandas.DataFrame`` objects, plus a bag of free text
artifacts (PDF prose, DOCX paragraphs, README files inside ZIPs, ...) that
the LLM planner can use as domain context.

Supported inputs
----------------
    tabular   : .csv .tsv .txt .parquet .xlsx .xls .xlsm
    semi      : .json .jsonl .ndjson .yaml .yml
    documents : .pdf (pdfplumber tables + text), .docx (tables + paragraphs)
    archives  : .zip (recursively unpacked, depth limited)
    remote    : any http(s) URL (content-type + extension sniffing)

Design notes
------------
* Every parser returns an :class:`IngestionResult`, so results compose with
  ``merge`` -- a ZIP of three CSVs and a PDF behaves exactly like four
  separate uploads.
* Nothing raises on a single bad member: failures are captured as warnings so
  a 200-file archive is not sunk by one corrupt entry.
* ``clean_dataframe`` is deliberately conservative. It fixes the things that
  break downstream sklearn code (duplicate/unnamed columns, all-null columns,
  numbers stored as strings, stray whitespace) and nothing else.
"""

from __future__ import annotations

import io
import json
import logging
import re
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

import pandas as pd
import yaml

LOGGER = logging.getLogger("autods.ingestion")

# --------------------------------------------------------------------------- #
# Optional dependencies -- degrade gracefully instead of exploding on import.
# --------------------------------------------------------------------------- #
try:  # pragma: no cover - environment dependent
    import pdfplumber
except Exception:  # pragma: no cover
    pdfplumber = None

try:  # pragma: no cover
    import docx as python_docx
except Exception:  # pragma: no cover
    python_docx = None

try:  # pragma: no cover
    import requests
except Exception:  # pragma: no cover
    requests = None


# --------------------------------------------------------------------------- #
# Configuration
# --------------------------------------------------------------------------- #
MAX_ZIP_MEMBERS = 250
MAX_ZIP_DEPTH = 3
MAX_MEMBER_BYTES = 256 * 1024 * 1024        # 256 MB per extracted member
MAX_PDF_PAGES = 60
MIN_TABLE_ROWS = 2                          # a 1-row "table" is usually noise
URL_TIMEOUT = 60

TABULAR_SUFFIXES = {".csv", ".tsv", ".txt", ".dat"}
EXCEL_SUFFIXES = {".xlsx", ".xls", ".xlsm", ".xlsb"}
JSON_SUFFIXES = {".json", ".jsonl", ".ndjson"}
YAML_SUFFIXES = {".yaml", ".yml"}
PARQUET_SUFFIXES = {".parquet", ".pq"}
SUPPORTED_SUFFIXES = (
    TABULAR_SUFFIXES
    | EXCEL_SUFFIXES
    | JSON_SUFFIXES
    | YAML_SUFFIXES
    | PARQUET_SUFFIXES
    | {".pdf", ".docx", ".zip"}
)

# Members we never bother opening inside an archive.
_ARCHIVE_SKIP = re.compile(
    r"(^__MACOSX/)|(/\._)|(^\._)|(\.DS_Store$)|(\.git/)|(\.ipynb_checkpoints/)"
)


# --------------------------------------------------------------------------- #
# Result containers
# --------------------------------------------------------------------------- #
@dataclass
class TextArtifact:
    """Free text pulled out of a document -- context for the planner, not data."""

    name: str
    source: str
    text: str

    def preview(self, limit: int = 2000) -> str:
        body = self.text.strip()
        return body if len(body) <= limit else body[:limit] + "\n... [truncated]"


@dataclass
class IngestionResult:
    """Everything recovered from one or more inputs."""

    frames: Dict[str, pd.DataFrame] = field(default_factory=dict)
    texts: List[TextArtifact] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)
    origins: Dict[str, str] = field(default_factory=dict)   # frame key -> source label

    # -- mutation ---------------------------------------------------------- #
    def add_frame(self, name: str, frame: pd.DataFrame, source: str = "") -> Optional[str]:
        """Register a frame under a unique, filesystem-safe key."""
        if frame is None or frame.empty or frame.shape[1] == 0:
            return None
        key = unique_key(slugify(name), self.frames)
        self.frames[key] = frame
        self.origins[key] = source or name
        return key

    def add_text(self, name: str, text: str, source: str = "") -> None:
        if text and text.strip():
            self.texts.append(TextArtifact(name=name, source=source or name, text=text))

    def warn(self, message: str) -> None:
        LOGGER.warning(message)
        self.warnings.append(message)

    def merge(self, other: "IngestionResult") -> "IngestionResult":
        for key, frame in other.frames.items():
            new_key = self.add_frame(key, frame, other.origins.get(key, ""))
            if new_key is None:
                self.warn(f"Dropped empty table '{key}'.")
        self.texts.extend(other.texts)
        self.warnings.extend(other.warnings)
        return self

    # -- reporting --------------------------------------------------------- #
    @property
    def is_empty(self) -> bool:
        return not self.frames

    def summary(self) -> pd.DataFrame:
        rows = []
        for key, frame in self.frames.items():
            missing = float(frame.isna().mean().mean() * 100) if frame.size else 0.0
            rows.append(
                {
                    "table": key,
                    "source": self.origins.get(key, ""),
                    "rows": len(frame),
                    "columns": frame.shape[1],
                    "numeric_cols": int(frame.select_dtypes("number").shape[1]),
                    "missing_%": round(missing, 2),
                    "memory_mb": round(frame.memory_usage(deep=True).sum() / 1e6, 3),
                }
            )
        return pd.DataFrame(rows)

    def primary_key(self) -> Optional[str]:
        """The frame most likely to be *the* modelling dataset (widest x tallest)."""
        if not self.frames:
            return None
        return max(self.frames.items(), key=lambda kv: (kv[1].shape[0] * kv[1].shape[1]))[0]


# --------------------------------------------------------------------------- #
# Naming helpers
# --------------------------------------------------------------------------- #
def slugify(name: str) -> str:
    stem = Path(str(name)).name
    stem = re.sub(r"\.[A-Za-z0-9]{1,6}$", "", stem)          # drop extension
    stem = re.sub(r"[^0-9a-zA-Z]+", "_", stem).strip("_").lower()
    stem = re.sub(r"_+", "_", stem)
    if not stem:
        stem = "table"
    if stem[0].isdigit():
        stem = f"t_{stem}"
    return stem[:60]


def unique_key(base: str, existing: Dict[str, Any]) -> str:
    if base not in existing:
        return base
    i = 2
    while f"{base}_{i}" in existing:
        i += 1
    return f"{base}_{i}"


# --------------------------------------------------------------------------- #
# Frame cleaning
# --------------------------------------------------------------------------- #
def _dedupe_columns(columns: Sequence[Any]) -> List[str]:
    seen: Dict[str, int] = {}
    out: List[str] = []
    for i, col in enumerate(columns):
        name = str(col).strip()
        name = re.sub(r"\s+", " ", name)
        if not name or name.lower().startswith("unnamed:") or name.lower() == "nan":
            name = f"column_{i + 1}"
        if name in seen:
            seen[name] += 1
            name = f"{name}_{seen[name]}"
        else:
            seen[name] = 0
        out.append(name)
    return out


def _is_texty(series: pd.Series) -> bool:
    """True for object *and* the ``str`` dtype pandas 3.x uses by default."""
    if pd.api.types.is_numeric_dtype(series) or pd.api.types.is_datetime64_any_dtype(series):
        return False
    if pd.api.types.is_bool_dtype(series) or isinstance(series.dtype, pd.CategoricalDtype):
        return False
    return series.dtype == object or pd.api.types.is_string_dtype(series)


def _maybe_numeric(series: pd.Series) -> pd.Series:
    """Convert text columns that are really numbers (incl. '1,234', '45%', '$9')."""
    if not _is_texty(series):
        return series
    sample = series.dropna().astype(str).head(200)
    if sample.empty:
        return series
    cleaned_sample = sample.str.replace(r"[,\s$%]", "", regex=True)
    parsed = pd.to_numeric(cleaned_sample, errors="coerce")
    if parsed.notna().mean() < 0.9:
        return series
    full = series.astype(str).str.replace(r"[,\s$%]", "", regex=True)
    converted = pd.to_numeric(full, errors="coerce")
    if sample.str.contains("%").mean() > 0.5:
        converted = converted / 100.0
    return converted


def _maybe_datetime(series: pd.Series) -> pd.Series:
    if not _is_texty(series):
        return series
    sample = series.dropna().astype(str).head(200)
    if sample.empty or sample.str.len().mean() < 6:
        return series
    if not sample.str.contains(r"[-/:]").mean() > 0.7:
        return series
    try:
        parsed = pd.to_datetime(series, errors="coerce", format="mixed")
    except (TypeError, ValueError):
        try:
            parsed = pd.to_datetime(series, errors="coerce")
        except Exception:
            return series
    return parsed if parsed.notna().mean() >= 0.9 else series


def clean_dataframe(frame: pd.DataFrame, infer_types: bool = True) -> pd.DataFrame:
    """Conservative normalisation: safe to run on every table we ever produce."""
    if frame is None or frame.empty:
        return pd.DataFrame()

    df = frame.copy()
    if isinstance(df.columns, pd.MultiIndex):
        df.columns = [" ".join(str(p) for p in tup if str(p) != "nan").strip() for tup in df.columns]
    df.columns = _dedupe_columns(df.columns)

    # Strip whitespace from string cells and normalise empty strings to NA.
    for col in df.columns:
        if _is_texty(df[col]):
            df[col] = df[col].apply(lambda v: v.strip() if isinstance(v, str) else v)
            df[col] = df[col].replace({"": None, "NA": None, "N/A": None, "null": None, "NULL": None, "-": None})

    df = df.dropna(axis=0, how="all").dropna(axis=1, how="all")
    if df.empty:
        return pd.DataFrame()

    if infer_types:
        for col in df.columns:
            converted = _maybe_numeric(df[col])
            if converted is df[col]:
                converted = _maybe_datetime(df[col])
            df[col] = converted

    return df.reset_index(drop=True)


# --------------------------------------------------------------------------- #
# Structured / semi-structured parsers
# --------------------------------------------------------------------------- #
def _read_delimited(data: bytes, name: str, result: IngestionResult) -> None:
    text = _decode(data)
    try:
        sample = text[:64_000]
        import csv as _csv

        try:
            dialect = _csv.Sniffer().sniff(sample, delimiters=",;\t|")
            sep = dialect.delimiter
        except Exception:
            sep = "\t" if name.lower().endswith(".tsv") else ","
        df = pd.read_csv(io.StringIO(text), sep=sep, engine="python", on_bad_lines="skip")
        result.add_frame(name, clean_dataframe(df), name)
    except Exception as exc:  # last resort: keep it as text context
        result.warn(f"{name}: delimited parse failed ({exc}); stored as text.")
        result.add_text(name, text, name)


def _read_excel(data: bytes, name: str, result: IngestionResult) -> None:
    try:
        book = pd.read_excel(io.BytesIO(data), sheet_name=None)
    except Exception as exc:
        result.warn(f"{name}: Excel parse failed ({exc}).")
        return
    for sheet, df in book.items():
        cleaned = clean_dataframe(df)
        if cleaned.empty:
            continue
        result.add_frame(f"{slugify(name)}_{slugify(sheet)}", cleaned, f"{name}#{sheet}")


def _read_parquet(data: bytes, name: str, result: IngestionResult) -> None:
    try:
        df = pd.read_parquet(io.BytesIO(data))
        result.add_frame(name, clean_dataframe(df, infer_types=False), name)
    except Exception as exc:
        result.warn(f"{name}: Parquet parse failed ({exc}). Is pyarrow installed?")


def _obj_to_frames(obj: Any, name: str, result: IngestionResult, depth: int = 0) -> None:
    """Walk a decoded JSON/YAML object and harvest every tabular structure."""
    if depth > 4 or obj is None:
        return

    if isinstance(obj, list):
        if obj and all(isinstance(item, dict) for item in obj):
            result.add_frame(name, clean_dataframe(pd.json_normalize(obj, max_level=2)), name)
        elif obj and all(not isinstance(item, (dict, list)) for item in obj):
            result.add_frame(name, clean_dataframe(pd.DataFrame({slugify(name): obj})), name)
        else:
            for i, item in enumerate(obj[:20]):
                _obj_to_frames(item, f"{name}_{i + 1}", result, depth + 1)
        return

    if isinstance(obj, dict):
        # dict-of-equal-length-lists -> a single table
        lists = {k: v for k, v in obj.items() if isinstance(v, list)}
        if lists and len(lists) == len(obj) and len({len(v) for v in lists.values()}) == 1:
            result.add_frame(name, clean_dataframe(pd.DataFrame(obj)), name)
            return
        # otherwise recurse into containers, and keep scalars as a key/value table
        scalars = {k: v for k, v in obj.items() if not isinstance(v, (dict, list))}
        for key, value in obj.items():
            if isinstance(value, (dict, list)):
                _obj_to_frames(value, f"{name}_{slugify(key)}", result, depth + 1)
        if scalars and depth == 0:
            kv = pd.DataFrame({"key": list(scalars.keys()), "value": list(scalars.values())})
            result.add_frame(f"{name}_metadata", clean_dataframe(kv), name)


def _read_json(data: bytes, name: str, result: IngestionResult) -> None:
    text = _decode(data)
    lowered = name.lower()
    if lowered.endswith((".jsonl", ".ndjson")):
        try:
            df = pd.read_json(io.StringIO(text), lines=True)
            result.add_frame(name, clean_dataframe(df), name)
            return
        except Exception as exc:
            result.warn(f"{name}: JSONL parse failed ({exc}); retrying as JSON.")
    try:
        obj = json.loads(text)
    except Exception as exc:
        result.warn(f"{name}: JSON parse failed ({exc}); stored as text.")
        result.add_text(name, text, name)
        return
    before = len(result.frames)
    _obj_to_frames(obj, slugify(name), result)
    if len(result.frames) == before:
        result.warn(f"{name}: no tabular structure found in JSON; stored as text.")
        result.add_text(name, text, name)


def _read_yaml(data: bytes, name: str, result: IngestionResult) -> None:
    text = _decode(data)
    try:
        docs = [d for d in yaml.safe_load_all(text) if d is not None]
    except Exception as exc:
        result.warn(f"{name}: YAML parse failed ({exc}); stored as text.")
        result.add_text(name, text, name)
        return
    before = len(result.frames)
    for i, doc in enumerate(docs):
        suffix = "" if len(docs) == 1 else f"_doc{i + 1}"
        _obj_to_frames(doc, f"{slugify(name)}{suffix}", result)
    if len(result.frames) == before:
        result.add_text(name, text, name)


# --------------------------------------------------------------------------- #
# Document parsers
# --------------------------------------------------------------------------- #
def _table_from_rows(rows: List[List[Any]]) -> Optional[pd.DataFrame]:
    rows = [r for r in rows if r and any(str(c).strip() for c in r if c is not None)]
    if len(rows) < MIN_TABLE_ROWS:
        return None
    header, body = rows[0], rows[1:]
    width = max(len(r) for r in rows)
    header = list(header) + [None] * (width - len(header))
    body = [list(r) + [None] * (width - len(r)) for r in body]
    return clean_dataframe(pd.DataFrame(body, columns=_dedupe_columns(header)))


def _read_pdf(data: bytes, name: str, result: IngestionResult) -> None:
    if pdfplumber is None:
        result.warn(f"{name}: pdfplumber is not installed -- PDF skipped.")
        return
    text_chunks: List[str] = []
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            pages = pdf.pages[:MAX_PDF_PAGES]
            if len(pdf.pages) > MAX_PDF_PAGES:
                result.warn(f"{name}: only the first {MAX_PDF_PAGES} pages were parsed.")
            for pno, page in enumerate(pages, start=1):
                try:
                    page_text = page.extract_text() or ""
                except Exception:
                    page_text = ""
                if page_text.strip():
                    text_chunks.append(f"[page {pno}]\n{page_text}")
                try:
                    tables = page.extract_tables() or []
                except Exception:
                    tables = []
                for tno, raw in enumerate(tables, start=1):
                    frame = _table_from_rows(raw)
                    if frame is not None and not frame.empty:
                        result.add_frame(
                            f"{slugify(name)}_p{pno}_t{tno}", frame, f"{name} page {pno} table {tno}"
                        )
    except Exception as exc:
        result.warn(f"{name}: PDF parse failed ({exc}).")
        return
    result.add_text(name, "\n\n".join(text_chunks), name)


def _read_docx(data: bytes, name: str, result: IngestionResult) -> None:
    if python_docx is None:
        result.warn(f"{name}: python-docx is not installed -- DOCX skipped.")
        return
    try:
        document = python_docx.Document(io.BytesIO(data))
    except Exception as exc:
        result.warn(f"{name}: DOCX parse failed ({exc}).")
        return
    for tno, table in enumerate(document.tables, start=1):
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        frame = _table_from_rows(rows)
        if frame is not None and not frame.empty:
            result.add_frame(f"{slugify(name)}_t{tno}", frame, f"{name} table {tno}")
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
    result.add_text(name, "\n".join(paragraphs), name)


# --------------------------------------------------------------------------- #
# Archives
# --------------------------------------------------------------------------- #
def _read_zip(data: bytes, name: str, result: IngestionResult, depth: int = 0) -> None:
    if depth >= MAX_ZIP_DEPTH:
        result.warn(f"{name}: maximum archive nesting depth reached.")
        return
    try:
        archive = zipfile.ZipFile(io.BytesIO(data))
    except Exception as exc:
        result.warn(f"{name}: not a readable ZIP ({exc}).")
        return

    members = [m for m in archive.infolist() if not m.is_dir() and not _ARCHIVE_SKIP.search(m.filename)]
    if len(members) > MAX_ZIP_MEMBERS:
        result.warn(f"{name}: {len(members)} members found; only the first {MAX_ZIP_MEMBERS} are read.")
        members = members[:MAX_ZIP_MEMBERS]

    for member in members:
        suffix = Path(member.filename).suffix.lower()
        if suffix not in SUPPORTED_SUFFIXES:
            continue
        if member.file_size > MAX_MEMBER_BYTES:
            result.warn(f"{name}:{member.filename} exceeds the per-file size limit; skipped.")
            continue
        try:
            payload = archive.read(member)
        except Exception as exc:
            result.warn(f"{name}:{member.filename} could not be extracted ({exc}).")
            continue
        label = f"{Path(name).stem}/{member.filename}"
        if suffix == ".zip":
            _read_zip(payload, label, result, depth + 1)
        else:
            result.merge(ingest_bytes(label, payload))


# --------------------------------------------------------------------------- #
# Entry points
# --------------------------------------------------------------------------- #
def _decode(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def ingest_bytes(name: str, data: bytes) -> IngestionResult:
    """Parse a single in-memory file. Never raises."""
    result = IngestionResult()
    suffix = Path(name).suffix.lower()
    try:
        if suffix in EXCEL_SUFFIXES:
            _read_excel(data, name, result)
        elif suffix in PARQUET_SUFFIXES:
            _read_parquet(data, name, result)
        elif suffix in JSON_SUFFIXES:
            _read_json(data, name, result)
        elif suffix in YAML_SUFFIXES:
            _read_yaml(data, name, result)
        elif suffix == ".pdf":
            _read_pdf(data, name, result)
        elif suffix == ".docx":
            _read_docx(data, name, result)
        elif suffix == ".zip":
            _read_zip(data, name, result)
        elif suffix in TABULAR_SUFFIXES:
            _read_delimited(data, name, result)
        else:
            # Unknown extension: sniff the payload.
            head = data[:4]
            if head[:2] == b"PK":
                _read_zip(data, name, result)
            elif head == b"%PDF":
                _read_pdf(data, name, result)
            else:
                _read_delimited(data, name, result)
    except Exception as exc:  # pragma: no cover - defensive
        result.warn(f"{name}: unhandled ingestion error ({exc}).")
    return result


def ingest_file(path: str | Path) -> IngestionResult:
    path = Path(path)
    if not path.exists():
        result = IngestionResult()
        result.warn(f"{path}: file not found.")
        return result
    return ingest_bytes(path.name, path.read_bytes())


def ingest_url(url: str) -> IngestionResult:
    """Download a remote resource and parse it by extension or content type."""
    result = IngestionResult()
    if requests is None:
        result.warn("The 'requests' package is required for URL ingestion.")
        return result
    try:
        response = requests.get(url, timeout=URL_TIMEOUT, headers={"User-Agent": "AutoDS-Agent/1.0"})
        response.raise_for_status()
    except Exception as exc:
        result.warn(f"{url}: download failed ({exc}).")
        return result

    name = Path(url.split("?")[0]).name or "remote_resource"
    if not Path(name).suffix:
        ctype = response.headers.get("content-type", "").split(";")[0].strip().lower()
        name += {
            "text/csv": ".csv",
            "application/json": ".json",
            "application/pdf": ".pdf",
            "application/zip": ".zip",
            "text/yaml": ".yaml",
            "application/x-yaml": ".yaml",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ".xlsx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
        }.get(ctype, ".csv")
    return result.merge(ingest_bytes(name, response.content))


def ingest_many(
    uploads: Optional[Iterable[Tuple[str, bytes]]] = None,
    urls: Optional[Iterable[str]] = None,
) -> IngestionResult:
    """Ingest a batch of uploads and/or URLs into one merged result."""
    combined = IngestionResult()
    for name, payload in uploads or []:
        combined.merge(ingest_bytes(name, payload))
    for url in urls or []:
        if url and url.strip():
            combined.merge(ingest_url(url.strip()))
    return combined


# --------------------------------------------------------------------------- #
# Profiling -- the structured view handed to the LLM planner
# --------------------------------------------------------------------------- #
def profile_frame(frame: pd.DataFrame, sample_rows: int = 5, max_columns: int = 120) -> Dict[str, Any]:
    n_rows = int(len(frame))
    columns: List[Dict[str, Any]] = []
    for col in list(frame.columns)[:max_columns]:
        series = frame[col]
        nunique = int(series.nunique(dropna=True))
        info: Dict[str, Any] = {
            "name": str(col),
            "dtype": str(series.dtype),
            "missing_pct": round(float(series.isna().mean() * 100), 2),
            "n_unique": nunique,
        }
        if pd.api.types.is_numeric_dtype(series) and series.notna().any():
            described = series.describe()
            info["stats"] = {
                k: (round(float(described[k]), 4) if pd.notna(described.get(k)) else None)
                for k in ("mean", "std", "min", "25%", "50%", "75%", "max")
                if k in described
            }
        elif nunique and nunique <= 25:
            counts = series.value_counts(dropna=True).head(10)
            info["top_values"] = {str(k): int(v) for k, v in counts.items()}
        columns.append(info)

    return {
        "n_rows": n_rows,
        "n_columns": int(frame.shape[1]),
        "truncated_columns": bool(frame.shape[1] > max_columns),
        "duplicate_rows": int(frame.duplicated().sum()) if n_rows <= 500_000 else None,
        "columns": columns,
        "sample": frame.head(sample_rows).to_dict(orient="records"),
    }


def profile_frames(
    frames: Dict[str, pd.DataFrame], sample_rows: int = 5
) -> Dict[str, Dict[str, Any]]:
    return {key: profile_frame(df, sample_rows=sample_rows) for key, df in frames.items()}


def json_safe(obj: Any) -> Any:
    """Recursively coerce numpy / pandas scalars so ``json.dumps`` never chokes."""
    import numpy as np

    if isinstance(obj, dict):
        return {str(k): json_safe(v) for k, v in obj.items()}
    if isinstance(obj, (list, tuple)):
        return [json_safe(v) for v in obj]
    if isinstance(obj, (np.integer,)):
        return int(obj)
    if isinstance(obj, (np.floating,)):
        value = float(obj)
        return value if value == value and abs(value) != float("inf") else None
    if isinstance(obj, (np.bool_,)):
        return bool(obj)
    if isinstance(obj, (pd.Timestamp,)):
        return obj.isoformat()
    if obj is pd.NaT or (isinstance(obj, float) and obj != obj):
        return None
    return obj


__all__ = [
    "IngestionResult",
    "TextArtifact",
    "clean_dataframe",
    "ingest_bytes",
    "ingest_file",
    "ingest_many",
    "ingest_url",
    "json_safe",
    "profile_frame",
    "profile_frames",
    "slugify",
]
